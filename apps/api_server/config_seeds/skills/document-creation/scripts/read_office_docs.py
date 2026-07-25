
import sys
import docx
import openpyxl

def read_docx(path):
    """Reads and prints text from a .docx file."""
    try:
        document = docx.Document(path)
        for para in document.paragraphs:
            print(para.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text, end='\\t')
                print()
    except Exception as e:
        print(f"Error reading docx file {path}: {e}", file=sys.stderr)

def read_xlsx(path):
    """Reads and prints data from an .xlsx file."""
    try:
        workbook = openpyxl.load_workbook(path)
        for sheet_name in workbook.sheetnames:
            print(f"--- Sheet: {sheet_name} ---")
            sheet = workbook[sheet_name]
            for row in sheet.iter_rows():
                print('\\t'.join([str(cell.value) if cell.value is not None else '' for cell in row]))
    except Exception as e:
        print(f"Error reading xlsx file {path}: {e}", file=sys.stderr)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python read_office_docs.py <path_to_file>", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]

    if file_path.endswith('.docx'):
        read_docx(file_path)
    elif file_path.endswith('.xlsx'):
        read_xlsx(file_path)
    else:
        print(f"Unsupported file type: {file_path}", file=sys.stderr)
        sys.exit(1)
